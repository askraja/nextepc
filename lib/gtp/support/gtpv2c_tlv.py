#
# Copyright (c) 2017, CellWire Group
# All rights reserved.
#
# Redistribution and use in source and binary forms, with or without
# modification, are permitted provided that the following conditions are met:
#
# 1. Redistributions of source code must retain the above copyright notice, this
#    list of conditions and the following disclaimer.
# 2. Redistributions in binary form must reproduce the above copyright notice,
#    this list of conditions and the following disclaimer in the documentation
#    and/or other materials provided with the distribution.
#
# THIS SOFTWARE IS PROVIDED BY THE COPYRIGHT HOLDERS AND CONTRIBUTORS "AS IS" AND
# ANY EXPRESS OR IMPLIED WARRANTIES, INCLUDING, BUT NOT LIMITED TO, THE IMPLIED
# WARRANTIES OF MERCHANTABILITY AND FITNESS FOR A PARTICULAR PURPOSE ARE
# DISCLAIMED. IN NO EVENT SHALL THE COPYRIGHT OWNER OR CONTRIBUTORS BE LIABLE FOR
# ANY DIRECT, INDIRECT, INCIDENTAL, SPECIAL, EXEMPLARY, OR CONSEQUENTIAL DAMAGES
# (INCLUDING, BUT NOT LIMITED TO, PROCUREMENT OF SUBSTITUTE GOODS OR SERVICES;
# LOSS OF USE, DATA, OR PROFITS; OR BUSINESS INTERRUPTION) HOWEVER CAUSED AND
# ON ANY THEORY OF LIABILITY, WHETHER IN CONTRACT, STRICT LIABILITY, OR TORT
# (INCLUDING NEGLIGENCE OR OTHERWISE) ARISING IN ANY WAY OUT OF THE USE OF THIS
# SOFTWARE, EVEN IF ADVISED OF THE POSSIBILITY OF SUCH DAMAGE.
#
from docx import Document
import re, os, sys, string
import datetime
import getopt
import getpass

version = "0.1.0"

msg_list = {}
type_list = {}
group_list = {}

verbosity = 0
filename = ""
outdir = './'
cachedir = './cache/'

FAIL = '\033[91m'
INFO = '\033[93m'
ENDC = '\033[0m'

def d_print(string):
    if verbosity > 0:
        sys.stdout.write(string)

def d_info(string):
    sys.stdout.write(INFO + string + ENDC + "\n")

def d_error(string):
    sys.stderr.write(FAIL + string + ENDC + "\n")
    sys.exit(0)

def write_file(f, string):
    f.write(string)
    d_print(string)

def output_header_to_file(f):
    now = datetime.datetime.now()
    f.write("""/*
 * Copyright (c) 2017, CellWire Group
 * All rights reserved.
 *
 * Redistribution and use in source and binary forms, with or without
 * modification, are permitted provided that the following conditions are met:
 * 
 * 1. Redistributions of source code must retain the above copyright notice, this
 *    list of conditions and the following disclaimer.
 * 2. Redistributions in binary form must reproduce the above copyright notice,
 *    this list of conditions and the following disclaimer in the documentation
 *    and/or other materials provided with the distribution.

 * THIS SOFTWARE IS PROVIDED BY THE COPYRIGHT HOLDERS AND CONTRIBUTORS "AS IS" AND
 * ANY EXPRESS OR IMPLIED WARRANTIES, INCLUDING, BUT NOT LIMITED TO, THE IMPLIED
 * WARRANTIES OF MERCHANTABILITY AND FITNESS FOR A PARTICULAR PURPOSE ARE
 * DISCLAIMED. IN NO EVENT SHALL THE COPYRIGHT OWNER OR CONTRIBUTORS BE LIABLE FOR
 * ANY DIRECT, INDIRECT, INCIDENTAL, SPECIAL, EXEMPLARY, OR CONSEQUENTIAL DAMAGES
 * (INCLUDING, BUT NOT LIMITED TO, PROCUREMENT OF SUBSTITUTE GOODS OR SERVICES;
 * LOSS OF USE, DATA, OR PROFITS; OR BUSINESS INTERRUPTION) HOWEVER CAUSED AND
 * ON ANY THEORY OF LIABILITY, WHETHER IN CONTRACT, STRICT LIABILITY, OR TORT
 * (INCLUDING NEGLIGENCE OR OTHERWISE) ARISING IN ANY WAY OUT OF THE USE OF THIS
 * SOFTWARE, EVEN IF ADVISED OF THE POSSIBILITY OF SUCH DAMAGE.
 */

""")
    f.write("/*******************************************************************************\n")
    f.write(" * This file had been created by gtpv2c_tlv.py script v%s\n" % (version))
    f.write(" * Please do not modify this file but regenerate it via script.\n")
    f.write(" * Created on: %s by %s\n * from %s\n" % (str(now), getpass.getuser(), filename))
    f.write(" ******************************************************************************/\n\n")

def usage():
    print "Python adding prefix for asn1 v%s" % (version)
    print "Usage: python asn1prefix.py [options]"
    print "Available options:"
    print "-d        Enable script debug"
    print "-f [file] Input file to parse"
    print "-o [dir]  Output files to given directory"
    print "-h        Print this help and return"

def v_upper(v):
    return re.sub('3GPP', '', re.sub('\'', '_', re.sub('/', '_', re.sub('-', '_', re.sub(' ', '_', v)))).upper())

def v_lower(v):
    return re.sub('3gpp', '', re.sub('\'', '_', re.sub('/', '_', re.sub('-', '_', re.sub(' ', '_', v)))).lower())

def get_cells(cells):
    instance = cells[4].text.encode('ascii', 'ignore')
    if instance.isdigit() is not True:
        return None
    ie_type = re.sub('\s*\n*\s*\(NOTE.*\)*', '', cells[3].text.encode('ascii', 'ignore'))
    if ie_type.find('LDN') != -1:
        ie_type = 'LDN'
    elif ie_type.find('APCO') != -1:
        ie_type = 'APCO'
    elif ie_type.find('Charging Id') != -1:
        ie_type = 'Charging ID'
    if ie_type not in type_list.keys():
        assert False, "Unknown IE type : [" \
                + cells[3].text + "]" + "(" + ie_type + ")"
    presence = cells[1].text.encode('ascii', 'ignore')
    ie_value = re.sub('\s*\n*\s*\([^\)]*\)*', '', cells[0].text).encode('ascii', 'ignore')
    comment = cells[2].text.encode('ascii', 'ignore')
    comment = re.sub('\n|\"|\'|\\\\', '', comment);

    if int(instance) > int(type_list[ie_type]["max_instance"]):
        type_list[ie_type]["max_instance"] = instance
        write_file(f, "type_list[\"" + ie_type + "\"][\"max_instance\"] = \"" + instance + "\"\n")

    return { "ie_type" : ie_type, "ie_value" : ie_value, "presence" : presence, "instance" : instance, "comment" : comment }

def write_cells_to_file(name, cells):
    write_file(f, name + ".append({ \"ie_type\" : \"" + cells["ie_type"] + \
        "\", \"ie_value\" : \"" + cells["ie_value"] + \
        "\", \"presence\" : \"" + cells["presence"] + \
        "\", \"instance\" : \"" + cells["instance"] + \
        "\", \"comment\" : \"" + cells["comment"] + "\"})\n")

try:
    opts, args = getopt.getopt(sys.argv[1:], "df:ho:c:", ["debug", "file", "help", "output", "cache"])
except getopt.GetoptError as err:
    # print help information and exit:
    usage()
    sys.exit(2)

for o, a in opts:
    if o in ("-d", "--debug"):
        verbosity = 1
    if o in ("-f", "--file"):
        filename = a
    if o in ("-o", "--output"):
        outdir = a
        if outdir.rfind('/') != len(outdir):
            outdir += '/'
    if o in ("-c", "--cache"):
        cache = a
        if cachedir.rfind('/') != len(cachedir):
            cachedir += '/'
    if o in ("-h", "--help"):
        usage()
        sys.exit(2)

if os.path.isfile(filename) and os.access(filename, os.R_OK):
    file = open(filename, 'r') 
else:
    d_error("Cannot find file : " + filename)

d_info("[Message List]")
cachefile = cachedir + 'gtpv2c_msg_list.py'
if os.path.isfile(cachefile) and os.access(cachefile, os.R_OK):
    execfile(cachefile)
    print "Read from " + cachefile
else:
    document = Document(filename)
    f = open(cachefile, 'w') 

    msg_table = ""
    for i, table in enumerate(document.tables):
        cell = table.rows[0].cells[0]
        if cell.text.find('Message Type value') != -1:
            msg_table = table
            d_print("Table Index = %d\n" % i)

    for row in msg_table.rows[2:-4]:
        key = row.cells[1].text.encode('ascii', 'ignore')
        type = row.cells[0].text.encode('ascii', 'ignore')
        if type.isdigit() is False:
            continue
        if int(type) in range(128, 160):
            continue
        if int(type) in range(231, 240):
            continue
        if key.find('Reserved') != -1:
            continue
        key = re.sub('\s*\n*\s*\([^\)]*\)*', '', key)
        msg_list[key] = { "type": type }
        write_file(f, "msg_list[\"" + key + "\"] = { \"type\" : \"" + type + "\" }\n")
    f.close()

d_info("[IE Type List]")
cachefile = cachedir + 'gtpv2c_type_list.py'
if os.path.isfile(cachefile) and os.access(cachefile, os.R_OK):
    execfile(cachefile)
    print "Read from " + cachefile
else:
    document = Document(filename)
    f = open(cachefile, 'w') 

    ie_table = ""
    for i, table in enumerate(document.tables):
        cell = table.rows[0].cells[0]
        if cell.text.find('IE Type value') != -1:
            ie_table = table
            d_print("Table Index = %d\n" % i)

    for row in ie_table.rows[1:-5]:
        key = row.cells[1].text.encode('ascii', 'ignore')
        if key.find('Reserved') != -1:
            continue
        if key.find('MM Context') != -1:
            continue
        elif key.find('Recovery') != -1:
            key = 'Recovery'
        elif key.find('Trusted WLAN Mode Indication') != -1:
            key = 'TWMI'
        elif key.find('LDN') != -1:
            key = 'LDN'
        elif key.find('APCO') != -1:
            key = 'APCO'
        elif key.find('Remote UE IP information') != -1:
            key = 'Remote UE IP Information'
        else:
            key = re.sub('.*\(', '', row.cells[1].text.encode('ascii', 'ignore'))
            key = re.sub('\)', '', key)
            key = re.sub('\s*$', '', key)
        type = row.cells[0].text.encode('ascii', 'ignore')
        type_list[key] = { "type": type , "max_instance" : "0" }
        write_file(f, "type_list[\"" + key + "\"] = { \"type\" : \"" + type)
        write_file(f, "\", \"max_instance\" : \"0\" }\n")
    f.close()
type_list['MM Context'] = { "type": "107", "max_instance" : "0" }

d_info("[Group IE List]")
cachefile = cachedir + 'gtpv2c_group_list.py'
if os.path.isfile(cachefile) and os.access(cachefile, os.R_OK):
    execfile(cachefile)
    print "Read from " + cachefile
else:
    document = Document(filename)
    f = open(cachefile, 'w') 

    for i, table in enumerate(document.tables):
        if table.rows[0].cells[0].text.find('Octet') != -1 and \
            table.rows[0].cells[2].text.find('IE Type') != -1:
            d_print("Table Index = %d\n" % i)

            row = table.rows[0];

            if len(re.findall('\d+', row.cells[2].text)) == 0:
                continue;
            ie_type = re.findall('\d+', row.cells[2].text)[0].encode('ascii', 'ignore')
            ie_name = re.sub('\s*IE Type.*', '', row.cells[2].text.encode('ascii', 'ignore'))

            if ie_name not in group_list.keys():
                group = []
                write_file(f, "group = []\n")
                for row in table.rows[4:]:
                    cells = get_cells(row.cells)
                    if cells is None:
                        continue

                    group_is_added = True
                    for ie in group:
                        if (cells["ie_type"], cells["instance"]) == (ie["ie_type"], ie["instance"]):
                            group_is_added = False
                    if group_is_added is True:
                        group.append(cells)
                        write_cells_to_file("group", cells)

                group_list[ie_name] = { "type" : ie_type, "group" : group }
                write_file(f, "group_list[\"" + ie_name + "\"] = { \"type\" : \"" + ie_type + "\", \"group\" : group }\n")
            else:
                group_list_is_added = False
                added_group = group_list[ie_name]["group"]
                write_file(f, "added_group = group_list[\"" + ie_name + "\"][\"group\"]\n")
                for row in table.rows[4:]:
                    cells = get_cells(row.cells)
                    if cells is None:
                        continue

                    group_is_added = True
                    for ie in group_list[ie_name]["group"]:
                        if (cells["ie_type"], cells["instance"]) == (ie["ie_type"], ie["instance"]):
                            group_is_added = False
                    for ie in group:
                        if (cells["ie_type"], cells["instance"]) == (ie["ie_type"], ie["instance"]):
                            group_is_added = False
                    if group_is_added is True:
                        added_group.append(cells)
                        write_cells_to_file("added_group", cells)
                        group_list_is_added = True
                if group_list_is_added is True:
                    group_list[ie_name] = { "type" : ie_type, "group" : added_group }
                    write_file(f, "group_list[\"" + ie_name + "\"] = { \"type\" : \"" + ie_type + "\", \"group\" : added_group }\n")
    f.close()

msg_list["Echo Request"]["table"] = 6
msg_list["Echo Response"]["table"] = 7
msg_list["Create Session Request"]["table"] = 8

for key in msg_list.keys():
    if "table" in msg_list[key].keys():
        d_info("[" + key + "]")
        cachefile = cachedir + "gtpv2c_msg_" + msg_list[key]["type"] + ".py"
        if os.path.isfile(cachefile) and os.access(cachefile, os.R_OK):
            execfile(cachefile)
            print "Read from " + cachefile
        else:
            document = Document(filename)
            f = open(cachefile, 'w') 

            ies = []
            write_file(f, "ies = []\n")
            table = document.tables[msg_list[key]["table"]]
            for row in table.rows[1:]:
                cells = get_cells(row.cells)
                if cells is None:
                    continue

                ies_is_added = True
                for ie in ies:
                    if (cells["ie_type"], cells["instance"]) == (ie["ie_type"], ie["instance"]):
                        ies_is_added = False
                if ies_is_added is True:
                    ies.append(cells)
                    write_cells_to_file("ies", cells)
            msg_list[key]["ies"] = ies
            write_file(f, "msg_list[key][\"ies\"] = ies\n")
            f.close()

f = open(outdir + 'gtpv2c_tlv.h', 'w')
output_header_to_file(f)
f.write("""#ifndef __GTPV2C_TLV_H__
#define __GTPV2C_TLV_H__

#include "core_tlv_msg.h"

#ifdef __cplusplus
extern "C" {
#endif /* __cplusplus */

""")

tmp = [(k, v["type"]) for k, v in msg_list.items()]
sorted_msg_list = sorted(tmp, key=lambda tup: int(tup[1]))
for (k, v) in sorted_msg_list:
    write_file(f, "#define GTPV2C_MSG_" + v_upper(k) + "_TYPE " + v + "\n")
write_file(f, "\n")

tmp = [(k, v["type"]) for k, v in type_list.items()]
sorted_type_list = sorted(tmp, key=lambda tup: int(tup[1]))
for (k, v) in sorted_type_list:
    write_file(f, "#define GTPV2C_IE_" + v_upper(k) + "_TYPE " + v + "\n")
write_file(f, "\n")

write_file(f, "/* Infomration Element TLV Descriptor */\n")
for (k, v) in sorted_type_list:
    if k not in group_list.keys():
        for instance in range(0, int(type_list[k]["max_instance"])+1):
            write_file(f, "extern tlv_desc_t gtpv2c_desc_" + v_lower(k))
            write_file(f, "_" + str(instance) + ";\n")
write_file(f, "\n")

tmp = [(k, v["type"]) for k, v in group_list.items()]
sorted_group_list = sorted(tmp, key=lambda tup: int(tup[1]))

write_file(f, "/* Group Infomration Element TLV Descriptor */\n")
for (k, v) in sorted_group_list:
    for instance in range(0, int(type_list[k]["max_instance"])+1):
        write_file(f, "extern tlv_desc_t gtpv2c_desc_" + v_lower(k))
        write_file(f, "_" + str(instance) + ";\n")
write_file(f, "\n")

write_file(f, "/* Message Descriptor */\n")
for (k, v) in sorted_msg_list:
    write_file(f, "extern tlv_desc_t gtpv2c_desc_" + v_lower(k) + ";\n")
write_file(f, "\n")

write_file(f, "/* Structure for Infomration Element */\n")
for (k, v) in sorted_type_list:
    if k not in group_list.keys():
        write_file(f, "typedef tlv_octet_t gtpv2c_" + v_lower(k) + "_t;\n")
write_file(f, "\n")

write_file(f, "/* Structure for Group Infomration Element */\n")
for (k, v) in sorted_group_list:
    write_file(f, "typedef struct _gtpv2c_" + v_lower(k) + "_t {\n")
    write_file(f, "    tlv_header_t h;\n")
    for group in group_list[k]["group"]:
        write_file(f, "    gtpv2c_" + v_lower(group["ie_type"]) + "_t " + \
                v_lower(group["ie_value"]))
        if group["ie_type"] == "F-TEID":
            if group["ie_value"] == "S2b-U ePDG F-TEID":
                write_file(f, "_" + group["instance"] + ";")
            elif group["ie_value"] == "S2a-U TWAN F-TEID":
                write_file(f, "_" + group["instance"] + ";")
            else:
                write_file(f, ";")
            write_file(f, " /* Instance : " + group["instance"] + " */\n")
        else:
            write_file(f, ";\n")
    write_file(f, "} gtpv2c_" + v_lower(k) + "_t;\n")
    write_file(f, "\n")

write_file(f, "/* Structure for Message */\n")
for (k, v) in sorted_msg_list:
    write_file(f, "typedef struct _gtpv2c_" + v_lower(k) + "_t {\n")
    write_file(f, "    tlv_header_t h;\n")
    if "ies" in msg_list[k]:
        for ies in msg_list[k]["ies"]:
            write_file(f, "    gtpv2c_" + v_lower(ies["ie_type"]) + "_t " + \
                    v_lower(ies["ie_value"]) + ";\n")
    write_file(f, "} gtpv2c_" + v_lower(k) + "_t;\n")
    write_file(f, "\n")

f.write("""#ifdef __cplusplus
}
#endif /* __cplusplus */

#endif /* __GTPV2C_TLV_H__ */
""")
f.close()

f = open(outdir + 'gtpv2c_tlv.c', 'w')
output_header_to_file(f)
f.write("""#include "gtpv2c_tlv.h"

""")

for (k, v) in sorted_type_list:
    if k not in group_list.keys():
        for instance in range(0, int(type_list[k]["max_instance"])+1):
            write_file(f, "tlv_desc_t gtpv2c_desc_%s_%d =\n" % (v_lower(k), instance))
            write_file(f, "{\n")
            write_file(f, "    TLV_VAR_STR,\n")
            write_file(f, "    GTPV2C_IE_%s_TYPE,\n" % v_upper(k))
            write_file(f, "    0,\n")
            write_file(f, "    %d,\n" % instance)
            write_file(f, "    sizeof(gtpv2c_%s_t),\n" % v_lower(k))
            write_file(f, "    { NULL }\n")
            write_file(f, "};\n\n")

for (k, v) in sorted_group_list:
    for instance in range(0, int(type_list[k]["max_instance"])+1):
        write_file(f, "tlv_desc_t gtpv2c_desc_%s_%d =\n" % (v_lower(k), instance))
        write_file(f, "{\n")
        write_file(f, "    TLV_COMPOUND,\n")
        write_file(f, "    GTPV2C_IE_%s_TYPE,\n" % v_upper(k))
        write_file(f, "    0,\n")
        write_file(f, "    %d,\n" % instance)
        write_file(f, "    sizeof(gtpv2c_%s_t),\n" % v_lower(k))
        write_file(f, "    {\n")
        for group in group_list[k]["group"]:
                write_file(f, "        &gtpv2c_desc_%s_%s,\n" % (v_lower(group["ie_type"]), v_lower(group["instance"])))
        write_file(f, "        NULL,\n")
        write_file(f, "    }\n")
        write_file(f, "};\n\n")

for (k, v) in sorted_msg_list:
    write_file(f, "tlv_desc_t gtpv2c_desc_%s =\n" % v_lower(k))
    write_file(f, "{\n")
    write_file(f, "    TLV_MESSAGE, 0, 0, 0, 0, {\n")
    if "ies" in msg_list[k]:
        for ies in msg_list[k]["ies"]:
                write_file(f, "        &gtpv2c_desc_%s_%s,\n" % (v_lower(ies["ie_type"]), v_lower(ies["instance"])))
    write_file(f, "    NULL,\n")
    write_file(f, "}};\n\n")

write_file(f, "\n")

f.close()
